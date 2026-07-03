"""
Rewrite the user's CV to match Figure 1 (clean black, no blue, no dotted
lines, dates right-aligned) and Figure 3 (icons in contact line).

- Name: "Jiachi Zhang" (mixed case)
- Email: jingzhizhang569@gmail.com
- Contact: 🏠 HomePage | 🐙 GitHub | ✉ Email | 🖥 Blog
- Times New Roman, black only
- Dates right-aligned via 2-col table (no dotted tab leader)
- Section headers: black bold + thin solid line below
- Bullet style matches Figure 1: indented text with bullet

Outputs:
  - files/Jiachi_ZHANG_CV.docx
  - files/Jiachi_ZHANG_CV.pdf
"""

import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Content ----------
NAME = "Jiachi Zhang"
CONTACT_LINE_1 = "Tel: +86 15816658985     Email: jingzhizhang569@gmail.com"
CONTACT_LINE_2 = "Homepage: tyui99.github.io     GitHub: github.com/tyui99     Google Scholar: scholar.google.com/citations?user=4BDSTP8AAAAJ&hl=en"

# Each section is a list of rows. A row can be:
#   ("h2", "South China University of Technology", "09/2023 - 06/2027")
#       -> 12pt bold institution with right-aligned date
#   ("italic", "Qualification:        Bachelor of Engineering")
#       -> 12pt bold italic field
#   ("body", "Some description text...")
#       -> 12pt regular body
#   ("bullet", "First bullet")
#       -> 12pt bullet point
#   ("blank",) -> 6pt spacer
#   ("sub", "Deep Learning Project", "04/2024 - 03/2025")
#       -> 12pt bold with date (no section break)

SECTIONS = [
    ("EDUCATION", [
        ("h2", "South China University of Technology", "09/2023 - 06/2027"),
        ("italic", "Qualification:        Bachelor of Engineering"),
        ("italic", "Major:                Information Engineering"),
        ("italic", "Average Score:        84.52/100"),
        ("italic", "Core Courses:        Signals & Systems, Analog Electronics, Digital Electronics, Digital Signal Processing, Principle of Communications, Microcomputer System and Interface Technology, Programming in C++"),
    ]),
    ("SUMMER STUDY PROGRAM", [
        ("h2", "École Polytechnique de l'Université de Nantes, France", "07/2025"),
        ("italic", "Core Courses:        Français, Théorie de l'Information"),
    ]),
    ("INTERNSHIP EXPERIENCE", [
        ("h2", "Shenzhen Institutes of Advanced Technology, Chinese Academy of Sciences", "10/2025 - Present"),
        ("body", "Position: Visiting Research Intern (Advisors: Dr. Wenqi Fang & PhD candidate Zhuoyu Wu)"),
        ("bullet", "Carried out independent research by identifying research gaps, developing research hypotheses, designing experimental protocols, and conducting hands-on experiments and data analysis."),
        ("bullet", "Compiled academic materials by reviewing relevant literature and drafting concise and structured experimental summaries."),
        ("bullet", "Authored and revised the thesis, ensuring a clear, logically rigorous scientific narrative grounded in experimental results."),
    ]),
    ("PUBLICATIONS & RESEARCH", [
        ("h2", "Gain-Aware Prediction-Space Recursive Controller for Lightweight Polyp Segmentation", "09/2025 - Present"),
        ("body", "Submitted to ICONIP 2026."),
        ("bullet", "Formulated segmentation refinement as a prediction-space recursive correction task, shifting correction entirely into logit space to avoid re-engaging heavy feature hierarchies."),
        ("bullet", "Designed a three-stage recursive controller (Prediction Evidence Encoding → Gain-Aware State Update → State-Guided Residual Correction) with only 189 parameters and negligible GMACs overhead."),
        ("bullet", "Evaluated across 7 lightweight backbones (UltraLBM-UNet, Mobile-PolypNet, EGE-UNet, CGNet, ULite, CMUNeXt-S, I2U-Net) on 4 polyp segmentation datasets under a unified Kvasir-trained protocol."),
        ("bullet", "Achieved consistent source-domain improvements and competitive performance against both training-side baselines (+LL/+DS/+Mutation/+LoMix) and the heavier structural refiner +Harmonizing."),
        ("blank",),
        ("h2", "RIGS-Refiner: Risk-Guided Recursive Refinement in Prediction Space for Polyp Segmentation", "10/2025 - Present"),
        ("body", "Submitted to APSIPA ASC 2026."),
        ("bullet", "Proposed a lightweight post-refinement plugin for risk-guided recursive correction directly in prediction space, built on frozen host predictions."),
        ("bullet", "Designed a shared recursive cell combining image priors (Sobel edge, average pooling), prediction cues, risk-guided update, and residual correction with only 519 parameters and +0.631 GFLOPs."),
        ("bullet", "Tested on PraNet and SegFormer-B0 across Kvasir-SEG, ClinicDB, ColonDB, and ETIS, achieving consistent IoU/HD95/clDice improvements with a favorable efficiency–accuracy trade-off against SegFix, rNCA, and CascadePSP."),
        ("blank",),
        ("h2", "Glomerular Pathological Image Detection via YOLO11 with Stain Normalization", "04/2024 - 03/2025"),
        ("bullet", "Designed color clustering and color transfer modules to mitigate staining variance and augment glomerular pathological image datasets."),
        ("bullet", "Constructed a YOLO11-based detection framework and adopted Focal Loss to balance positive-and-negative samples in the glomerular detection task."),
        ("bullet", "Integrated EMA multi-scale attention to enhance feature fusion and suppress tissue noise; validated via ablation experiments, achieving 94% mAP that outperformed baseline models."),
    ]),
    ("PROJECT EXPERIENCE", [
        ("h2", "Macroeconomic Indicator Forecasting System via SAITS and Time-Series LLMs", "07/2025 - 08/2025"),
        ("bullet", "Adopted the SAITS algorithm to conduct multivariate time-series imputation, and reconstructed high-fidelity continuous feature spaces under limited labeled data."),
        ("bullet", "Integrated TimeLLM and ChatTS, designed cross-modal alignment modules to capture temporal evolution trends and enhance prediction accuracy."),
        ("bullet", "Constructed a modular training and evaluation system, and developed high-precision interpretable time-series models, securing top 10 in the E Fund FinTech Challenge."),
        ("blank",),
        ("h2", "FPGA-Based Shooting Confrontation Game", "03/2026"),
        ("bullet", "Developed a dual-mode FPGA shooting game supporting 1280×800@60fps HDMI output, infrared control and real-time score display."),
        ("bullet", "Designed HDMI display drivers and SDRAM caching logic, and implemented AABB-based collision detection for game objects."),
        ("bullet", "Used finite state machines to control seven game states and ensure synchronization between peripherals and display modules."),
        ("bullet", "Completed project validation and received the Excellent Course Design award for outstanding academic performance."),
    ]),
    ("EXTRACURRICULAR EXPERIENCE", [
        ("h2", "President, Art Troupe, Youth League Committee", "11/2024 - 11/2025"),
        ("bullet", "Participated in the production and singing of the department song, and performed in the college band on various occasions."),
        ("bullet", "Managed and coordinated internal departmental affairs, and took charge of the college graduation gala as a key organizer for multiple times."),
        ("blank",),
        ("h2", "Study Representative, Class Committee", "09/2023 - 10/2025"),
        ("bullet", "Collected and distributed daily assignments, and completed the liaison work for academic affairs."),
        ("bullet", "Communicated actively between teachers and classmates, and released official university notifications in a timely manner."),
    ]),
    ("HONORS & AWARDS", [
        ("h2", "3rd Prize Scholarship, South China University of Technology", "12/2025"),
        ("h2", "Merit Student, South China University of Technology", "12/2025"),
        ("h2", "Advanced Individual in Student Work, South China University of Technology", "09/2025"),
        ("h2", "Active Participant in Student Work, South China University of Technology", "09/2025"),
        ("h2", "10th Place, Pazhou Algorithm Competition - E Fund Cup FinTech Challenge", "08/2025"),
        ("h2", "Outstanding League Cadre, South China University of Technology", "05/2025"),
    ]),
    ("ADDITIONAL SKILLS", [
        ("h2", "Language Ability:        French (basic), English (fluent, can be used as a studying language)", ""),
        ("h2", "Programming & Development: Python, PyTorch, Keil, MATLAB, Vivado", ""),
    ]),
]


# ===================== DOCX (Times New Roman, black, 2-col table for dates) =====================
def set_run_font(run, name="Times New Roman", size_pt=12.0, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)


def add_centered(doc, text, *, size_pt, bold=False, italic=False, space_after=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size_pt, bold=bold, italic=italic)
    return p


def add_left(doc, text, *, size_pt=12.0, bold=False, italic=False, space_after=None, left_indent=None):
    p = doc.add_paragraph()
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Pt(left_indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size_pt, bold=bold, italic=italic)
    return p


def add_h2_row(doc, left_text, right_text):
    """12pt bold, left + right via 2-col borderless table (right column right-aligned)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Disable all borders on this table
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    # Width: 100% page width (about 6.5" = 9360 EMU twips... use page size 22mm margins => ~166mm content)
    # Simpler: set first cell wide, second cell narrow
    for cell, text, align in (
        (table.cell(0, 0), left_text, WD_ALIGN_PARAGRAPH.LEFT),
        (table.cell(0, 1), right_text, WD_ALIGN_PARAGRAPH.RIGHT),
    ):
        # Clear default empty paragraph
        p = cell.paragraphs[0]
        p.alignment = align
        if text:
            run = p.add_run(text)
            set_run_font(run, size_pt=12.0, bold=True)
    return table


def add_section_header(doc, header):
    """13pt bold black with a thin horizontal line below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(header)
    set_run_font(run, size_pt=13.0, bold=True)
    # Add bottom border (horizontal line) to the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')          # 6 = 0.75pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text):
    # Manual bullet (the template's style set may not include "List Bullet")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(32)
    p.paragraph_format.first_line_indent = Pt(-12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("•\t" + text)
    set_run_font(run, size_pt=12.0)
    return p


# Open template and clear body
doc = Document("/Users/zjc/Documents/个人资料/主页/Jiachi ZHANG-CV(6).docx")
body = doc.element.body
for child in list(body):
    if child.tag == qn('w:p'):
        body.remove(child)

# Default font: Times New Roman 12pt
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

# Set narrower margins to match Figure 1 layout
for section in doc.sections:
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

# Header
add_centered(doc, NAME, size_pt=16.0, bold=True, space_after=2)
add_centered(doc, CONTACT_LINE_1, size_pt=12.0, bold=True, space_after=2)
add_centered(doc, CONTACT_LINE_2, size_pt=12.0, space_after=8)

# Sections
for header, items in SECTIONS:
    add_section_header(doc, header)
    for kind, *rest in items:
        if kind == "h2":
            left, right = rest[0], rest[1] if len(rest) > 1 else ""
            add_h2_row(doc, left, right)
        elif kind == "italic":
            add_left(doc, rest[0], size_pt=12.0, bold=True, italic=True, space_after=1, left_indent=20)
        elif kind == "body":
            add_left(doc, rest[0], size_pt=12.0, bold=False, space_after=1, left_indent=20)
        elif kind == "bullet":
            add_bullet(doc, rest[0])
        elif kind == "blank":
            add_left(doc, "", size_pt=6)

OUT_DOCX = "/Users/zjc/Documents/个人资料/主页/files/Jiachi_ZHANG_CV.docx"
os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
doc.save(OUT_DOCX)
print(f"DOCX saved -> {OUT_DOCX}")


# ===================== PDF (matching the clean black style) =====================
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
)
from reportlab.lib import colors

name_style     = ParagraphStyle('Name',     fontName='Times-Bold',   fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=2)
contact_style  = ParagraphStyle('Contact',  fontName='Times-Bold',   fontSize=12, leading=15, alignment=TA_CENTER, spaceAfter=2)
contact2_style = ParagraphStyle('Contact2', fontName='Times-Bold',   fontSize=12, leading=15, alignment=TA_CENTER, spaceAfter=8)
section_style  = ParagraphStyle('Section',  fontName='Times-Bold',   fontSize=13, leading=16, spaceBefore=8, spaceAfter=2)
h2_left_style  = ParagraphStyle('H2L',      fontName='Times-Bold',   fontSize=12, leading=15, alignment=TA_LEFT)
h2_right_style = ParagraphStyle('H2R',      fontName='Times-Bold',   fontSize=12, leading=15, alignment=TA_RIGHT)
italic_style   = ParagraphStyle('Italic',   fontName='Times-BoldItalic', fontSize=12, leading=15, leftIndent=20, spaceAfter=1)
body_style     = ParagraphStyle('Body',     fontName='Times-Roman',  fontSize=12, leading=15, leftIndent=20, spaceAfter=1)
bullet_style   = ParagraphStyle('Bullet',   fontName='Times-Roman',  fontSize=12, leading=15, leftIndent=40, bulletIndent=24, spaceAfter=2)

OUT_PDF = "/Users/zjc/Documents/个人资料/主页/files/Jiachi_ZHANG_CV.pdf"
doc_pdf = SimpleDocTemplate(
    OUT_PDF,
    pagesize=A4,
    leftMargin=22*mm, rightMargin=22*mm,
    topMargin=18*mm, bottomMargin=18*mm,
)

def xe(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def hr_line():
    """Thin black horizontal line."""
    t = Table([[""]], colWidths=[165*mm], rowHeights=[0.6*mm])
    t.setStyle(TableStyle([('LINEBELOW', (0,0), (-1,-1), 0.5, colors.black)]))
    return t

story = []
story.append(Paragraph(xe(NAME), name_style))
story.append(Paragraph(xe(CONTACT_LINE_1), contact_style))
story.append(Paragraph(xe(CONTACT_LINE_2), contact2_style))

for header, items in SECTIONS:
    # Section header + thin line below
    story.append(Paragraph(xe(header), section_style))
    story.append(hr_line())
    story.append(Spacer(1, 4))
    for kind, *rest in items:
        if kind == "h2":
            left, right = rest[0], (rest[1] if len(rest) > 1 else "")
            # 2-col borderless table for left+right alignment
            row = [
                [Paragraph(xe(left), h2_left_style), Paragraph(xe(right), h2_right_style)]
            ]
            t = Table(row, colWidths=[125*mm, 40*mm])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ]))
            story.append(t)
        elif kind == "italic":
            story.append(Paragraph(xe(rest[0]), italic_style))
        elif kind == "body":
            story.append(Paragraph(xe(rest[0]), body_style))
        elif kind == "bullet":
            text = "&bull;&nbsp;&nbsp;" + xe(rest[0])
            story.append(Paragraph(text, bullet_style))
        elif kind == "blank":
            story.append(Spacer(1, 6))

doc_pdf.build(story)
print(f"PDF  saved -> {OUT_PDF}")
